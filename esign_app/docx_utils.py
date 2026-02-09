"""
DOCX Processing utilities for extracting text and adding signatures
"""
import os
from typing import Dict, List, Tuple
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from copy import deepcopy
import fitz  # PyMuPDF
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.utils import ImageReader
from reportlab.platypus import Paragraph, Spacer, Table, TableStyle, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from io import BytesIO
import PIL.Image


class DOCXProcessor:
    @staticmethod
    def get_docx_info(docx_path: str) -> Dict:
        """Get basic information about a DOCX file"""
        try:
            doc = Document(docx_path)
            # Estimate number of pages by simulating page division
            all_content = []
            for para in doc.paragraphs:
                if para.text.strip():
                    all_content.append(para.text)
            
            # Count pages based on character estimation
            char_per_page = 1000  # Same as in extract_text_from_docx
            num_pages = 1
            current_char_count = 0
            
            for para_text in all_content:
                para_char_count = len(para_text)
                if current_char_count + para_char_count > char_per_page:
                    num_pages += 1
                    current_char_count = para_char_count
                else:
                    current_char_count += para_char_count
            
            return {
                'num_paragraphs': len(doc.paragraphs),
                'num_tables': len(doc.tables),
                'num_pages': num_pages,  # Add page count for consistency with PDF
                'has_images': len([rel for rel in doc.part.rels.values() if "image" in rel.target_ref]) > 0,
                'extension': 'docx'
            }
        except Exception as e:
            print(f"Error getting DOCX info: {e}")
            return {'extension': 'docx', 'num_pages': 1}
    
    @staticmethod
    def extract_text_from_docx(docx_path: str) -> Dict[int, str]:
        """
        Extract text content from DOCX organized by simulated pages
        """
        try:
            doc = Document(docx_path)
            pages_dict = {}
            
            # Process paragraphs and organize them into pages
            all_content = []
            for para in doc.paragraphs:
                if para.text.strip():
                    all_content.append(para.text)
            
            # Estimate content per page based on character count (simulating pages)
            # Assuming roughly 1000 characters per page as a rough estimate
            char_per_page = 1000  # Adjust this value as needed
            current_char_count = 0
            page_num = 1
            current_page_content = []
            
            for para_text in all_content:
                para_char_count = len(para_text)
                
                # If adding this paragraph would exceed the page limit, start a new page
                if current_char_count + para_char_count > char_per_page and current_page_content:
                    pages_dict[page_num] = '\n'.join(current_page_content)
                    page_num += 1
                    current_page_content = [para_text]
                    current_char_count = para_char_count
                else:
                    current_page_content.append(para_text)
                    current_char_count += para_char_count
            
            # Add the last page if there's content
            if current_page_content:
                pages_dict[page_num] = '\n'.join(current_page_content)
            
            # If no pages were created, add the whole content as one page
            if not pages_dict and all_content:
                pages_dict[1] = '\n'.join(all_content)
            
            return pages_dict
        except Exception as e:
            print(f"Error extracting text from DOCX: {e}")
            return {1: "Error extracting document content"}
    
    @staticmethod
    def add_signature_to_docx_with_position(
        docx_path: str,
        signature_path: str,
        output_path: str,
        x: float,
        y: float,
        page_number: int = 1
    ) -> bool:
        """
        Add signature to DOCX at specific coordinates by converting to PDF first
        
        Args:
            docx_path: Path to input DOCX
            signature_path: Path to signature image
            output_path: Path to save signed DOCX (will be converted to PDF)
            x: X coordinate for signature placement
            y: Y coordinate for signature placement
            page_number: Page number to add signature
            
        Returns:
            True if successful, False otherwise
        """
        try:
            print(f"=== DOCX CUSTOM POSITIONING START ===")
            print(f"Input DOCX: {docx_path}")
            print(f"Signature image: {signature_path}")
            print(f"Output path: {output_path}")
            print(f"Position: x={x}, y={y}")
            
            # Check if signature file exists
            if not os.path.exists(signature_path):
                print(f"ERROR: Signature file not found at {signature_path}")
                return False
            
            # Convert DOCX to PDF first
            pdf_path = docx_path.replace('.docx', '_temp.pdf')
            if not DOCXProcessor.docx_to_pdf(docx_path, pdf_path):
                print("Failed to convert DOCX to PDF")
                return False
            
            # Add signature to PDF at specified position
            from .pdf_utils import PDFProcessor
            success = PDFProcessor.add_signature_to_pdf(
                input_pdf_path=pdf_path,
                output_pdf_path=output_path,
                signature_image_path=signature_path,
                page_number=page_number,
                x=x,
                y=y,
                label_text="Signature",
                signature_width=150,
                signature_height=50
            )
            
            # Clean up temporary PDF
            if os.path.exists(pdf_path):
                os.remove(pdf_path)
            
            print(f"=== DOCX CUSTOM POSITIONING END ===")
            return success
            
        except Exception as e:
            print(f"Error adding signature to DOCX with custom position: {e}")
            import traceback
            traceback.print_exc()
            return False
    
    @staticmethod
    def convert_docx_to_pdf(input_docx: str, output_pdf: str) -> bool:
        """
        Convert DOCX to PDF
        """
        return DOCXProcessor.docx_to_pdf(input_docx, output_pdf)
    
    @staticmethod
    def docx_to_pdf(docx_path: str, pdf_path: str) -> bool:
        """
        Convert DOCX to PDF using reportlab
        """
        try:
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import Image as ReportLabImage
            
            doc = Document(docx_path)
            
            # Create PDF document
            pdf_doc = SimpleDocTemplate(pdf_path, pagesize=letter)
            story = []
            styles = getSampleStyleSheet()
            
            # Process paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    p = Paragraph(para.text, styles['Normal'])
                    story.append(p)
                    story.append(Spacer(1, 12))
            
            # Process tables
            for table in doc.tables:
                # Convert table to simple text representation
                table_text = ""
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    table_text += row_text + "<br/>"
                if table_text:
                    p = Paragraph(table_text, styles['Normal'])
                    story.append(p)
                    story.append(Spacer(1, 12))
            
            # Build PDF
            pdf_doc.build(story)
            return True
            
        except Exception as e:
            print(f"Error converting DOCX to PDF: {e}")
            return False
    
    @staticmethod
    def add_signature_to_docx(
        docx_path: str,
        signature_path: str,
        output_path: str,
        x: float,
        y: float,
        page_number: int = 1
    ) -> bool:
        """
        Add signature to DOCX at specific coordinates (limited support)
        For DOCX, this will place the signature at the end of the document
        """
        # For DOCX, we'll add the signature at the end of the document
        # Since DOCX doesn't have traditional pages like PDFs, we'll append signatures
        return DOCXProcessor.add_signature_to_docx_with_alignment(
            docx_path, signature_path, output_path, 
            is_right_aligned=(x > 400),  # If x is in right half, align right
            page_number=page_number
        )

    @staticmethod
    def add_signature_to_docx_preserve_format(
        docx_path: str,
        signature_path: str,
        output_path: str,
        x: float,
        y: float,
        page_number: int = 1
    ) -> bool:
        """
        Add signature to DOCX while preserving original DOCX format
        Uses python-docx to add signature as image with proper positioning
        """
        try:
            print(f"=== DOCX PRESERVE FORMAT START ===")
            print(f"Input DOCX: {docx_path}")
            print(f"Signature image: {signature_path}")
            print(f"Output path: {output_path}")
            print(f"Position: x={x}, y={y}")
            
            # Check if signature file exists
            import os
            if not os.path.exists(signature_path):
                print(f"ERROR: Signature file not found at {signature_path}")
                # Create a placeholder text instead of the image
                doc = Document(docx_path)
                doc.add_paragraph()  # Add spacing
                paragraph = doc.add_paragraph()
                
                # If x > 400 pixels (roughly right half of page), align right
                is_right_aligned = x > 400
                if is_right_aligned:
                    from docx.enum.text import WD_ALIGN_PARAGRAPH
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    print("Setting right alignment for signature placeholder")
                else:
                    from docx.enum.text import WD_ALIGN_PARAGRAPH
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    print("Setting left alignment for signature placeholder")
                
                run_label = paragraph.add_run("Signature: ")
                run_placeholder = paragraph.add_run("[Signature image not available]")
                run_placeholder.bold = True
                
                # Add metadata
                metadata_para = doc.add_paragraph()
                from docx.shared import Pt
                metadata_run = metadata_para.add_run(f"\n[Signature placed at coordinates: x={x}, y={y}, page={page_number}]")
                metadata_run.font.size = Pt(1)
                
                doc.save(output_path)
                print(f"Created DOCX with signature placeholder")
                return True
            
            doc = Document(docx_path)
            
            # If x > 400 pixels (roughly right half of page), align right
            is_right_aligned = x > 400
            
            # Add signature at the end of document with proper alignment
            doc.add_paragraph()  # Add spacing
            
            # Add signature label and image in the same paragraph for proper alignment
            from docx.shared import Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            paragraph = doc.add_paragraph()
            
            if is_right_aligned:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                print("Setting right alignment for signature")
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                print("Setting left alignment for signature")
            
            run_label = paragraph.add_run("Signature: ")
            
            run_image = paragraph.add_run()
            run_image.add_picture(signature_path, width=Inches(2))
            
            # Add the position coordinates as hidden metadata for reference
            # This helps maintain the intended positioning info
            metadata_para = doc.add_paragraph()
            metadata_run = metadata_para.add_run(f"\n[Signature placed at coordinates: x={x}, y={y}, page={page_number}]")
            from docx.shared import Pt
            metadata_run.font.size = Pt(1)  # Very small font size to make it nearly invisible
            
            doc.save(output_path)
            
            print(f"=== DOCX PRESERVE FORMAT END ===")
            return True
            
        except Exception as e:
            print(f"Error adding signature to DOCX while preserving format: {e}")
            import traceback
            traceback.print_exc()
            return False

    @staticmethod
    def sign_docx_all_pages_footer(docx_path: str, sign_img_path: str, output_path: str, width_inches: float = 1.5) -> bool:
        """
        Place the signature image into the footer of every section in the DOCX.
        This is a simple, reliable way to show the same signature on every
        page (or section) of the document.
        """
        try:
            doc = Document(docx_path)
            for section in doc.sections:
                footer = section.footer
                # Ensure there is at least one paragraph in the footer
                if not footer.paragraphs:
                    para = footer.add_paragraph()
                else:
                    para = footer.paragraphs[0]

                run = para.add_run()
                run.add_picture(sign_img_path, width=Inches(width_inches))

            doc.save(output_path)
            return True
        except Exception as e:
            print(f"Error placing signature in footers: {e}")
            import traceback
            traceback.print_exc()
            return False

    @staticmethod
    def add_signature_to_docx_end(
        docx_path: str,
        signature_path: str,
        output_path: str,
        is_right_aligned: bool = False,
        width_inches: float = 2.0
    ) -> bool:
        """
        Append a signature to the end of a DOCX document.

        This places the image as a new paragraph at the end of the document
        and aligns it left or right depending on `is_right_aligned`.
        """
        try:
            if not os.path.exists(signature_path):
                print(f"ERROR: Signature file not found at {signature_path}")
                return False

            doc = Document(docx_path)
            doc.add_paragraph()  # spacing

            from docx.shared import Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            para = doc.add_paragraph()
            if is_right_aligned:
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT

            run_label = para.add_run("Signature: ")
            run_image = para.add_run()
            run_image.add_picture(signature_path, width=Inches(width_inches))

            # Small metadata to indicate placement
            meta = doc.add_paragraph()
            meta_run = meta.add_run(f"[Signature appended at end: aligned={'right' if is_right_aligned else 'left'}]")
            from docx.shared import Pt
            meta_run.font.size = Pt(1)

            doc.save(output_path)
            print(f"Appended signature to end of DOCX: {output_path}")
            return True
        except Exception as e:
            print(f"Error appending signature to DOCX end: {e}")
            import traceback
            traceback.print_exc()
            return False
    
    @staticmethod
    def add_multiple_signatures_to_docx(
        docx_path: str,
        signatures_data: List[Dict],
        output_path: str
    ) -> bool:
        """
        Add multiple signatures to DOCX file with guaranteed end-of-page placement
        
        Args:
            docx_path: Path to input DOCX
            signatures_data: List of signature data dicts
            output_path: Path to save signed DOCX
        
        Returns:
            True if successful, False otherwise
        """
        try:
            print(f"=== DOCX END-POSITION SIGNATURES START ===")
            print(f"Input DOCX: {docx_path}")
            print(f"Output path: {output_path}")
            print(f"Number of signatures to add: {len(signatures_data)}")
            
            # Ensure output directory exists
            output_dir = os.path.dirname(output_path)
            print(f"Output directory: {output_dir}")
            os.makedirs(output_dir, exist_ok=True)
            print(f"Directory exists after creation: {os.path.exists(output_dir)}")
            
            # Check if all signature files exist
            for sig_data in signatures_data:
                if not os.path.exists(sig_data['signature_image_path']):
                    print(f"ERROR: Signature file not found at {sig_data['signature_image_path']}")
                    return False
            
            # Load the original document
            doc = Document(docx_path)
            
            # Group signatures by page
            signatures_by_page = {}
            end_position_signatures = []
            
            for sig_data in signatures_data:
                page_num = sig_data.get('page', 1)
                is_end_position = sig_data.get('is_end_position', False)
                
                if is_end_position:
                    # Collect all end-position signatures
                    end_position_signatures.append(sig_data)
                    print(f"Found end-position signature for page {page_num}")
                else:
                    # Regular positioning
                    if page_num not in signatures_by_page:
                        signatures_by_page[page_num] = []
                    signatures_by_page[page_num].append(sig_data)
            
            print(f"Regular signatures by page: {signatures_by_page}")
            print(f"End-position signatures: {len(end_position_signatures)}")
            
            # For end-position signatures, use footer placement
            if end_position_signatures:
                print("Placing end-position signatures in footers...")
                success = DOCXProcessor._place_signatures_in_footers(doc, end_position_signatures)
                if not success:
                    print("Failed to place signatures in footers")
                    return False
            
            # Handle regular positioned signatures (if any)
            if signatures_by_page:
                print("Handling regular positioned signatures...")
                success = DOCXProcessor._place_regular_signatures(doc, signatures_by_page)
                if not success:
                    print("Failed to place regular signatures")
                    return False
            
            # Save the document
            doc.save(output_path)
            
            # Verify file was created
            file_created = os.path.exists(output_path)
            file_size = os.path.getsize(output_path) if file_created else 0
            print(f"File created: {file_created}, Size: {file_size} bytes")
            
            print(f"=== DOCX END-POSITION SIGNATURES END ===")
            return file_created
        
        except Exception as e:
            print(f"Error adding multiple signatures to DOCX: {e}")
            import traceback
            traceback.print_exc()
            return False
    
    @staticmethod
    def _place_signatures_in_footers(doc, signatures_data: List[Dict]) -> bool:
        """Place signatures in document footers for guaranteed end-of-page positioning"""
        try:
            from docx.shared import Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            # Process each signature for footer placement
            for i, sig_data in enumerate(signatures_data):
                signature_path = sig_data.get('signature_image_path')
                x = sig_data.get('x', 0)
                page_num = sig_data.get('page', 1)
                
                # Determine alignment
                align_field = sig_data.get('align')
                if align_field is not None:
                    is_right_aligned = str(align_field).lower() == 'right'
                else:
                    is_right_aligned = x > 400
                
                print(f"Placing end-position signature {i+1} in footer (page {page_num}, alignment: {'right' if is_right_aligned else 'left'})")
                
                # For simplicity, place in all sections' footers
                # In a real implementation, you'd target specific sections/pages
                for section in doc.sections:
                    footer = section.footer
                    
                    # Clear existing content
                    for paragraph in footer.paragraphs[:]:
                        p_element = paragraph._element
                        p_element.getparent().remove(p_element)
                        paragraph._p = paragraph._element = None
                    
                    # Add signature paragraph
                    sig_para = footer.add_paragraph()
                    
                    # Set alignment
                    if is_right_aligned:
                        sig_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    else:
                        sig_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # Add signature label
                    run_label = sig_para.add_run("Signature")
                    run_label.bold = True
                    
                    # Add signature image
                    run_image = sig_para.add_run()
                    try:
                        image_width = Inches(2.0)
                        if is_right_aligned:
                            image_width = Inches(1.8)
                        run_image.add_picture(signature_path, width=image_width)
                    except Exception as e:
                        print(f"Error adding signature image to footer: {e}")
                        run_image.add_run("[SIGNATURE IMAGE]")
                
                print(f"Successfully placed end-position signature in footers")
            
            return True
            
        except Exception as e:
            print(f"Error placing signatures in footers: {e}")
            import traceback
            traceback.print_exc()
            return False
    
    @staticmethod
    def _place_regular_signatures(doc, signatures_by_page: Dict) -> bool:
        """Place regular (non-end-position) signatures in document content"""
        try:
            from docx.shared import Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            # Simple approach: place at the end of document for regular signatures too
            # This ensures they appear at the end regardless of original positioning
            for page_num, sig_list in signatures_by_page.items():
                for sig_data in sig_list:
                    signature_path = sig_data.get('signature_image_path')
                    x = sig_data.get('x', 0)
                    
                    # Determine alignment
                    align_field = sig_data.get('align')
                    if align_field is not None:
                        is_right_aligned = str(align_field).lower() == 'right'
                    else:
                        is_right_aligned = x > 400
                    
                    print(f"Placing regular signature at end of document (page {page_num})")
                    
                    # Add spacing to push to end
                    for _ in range(3):
                        spacing_para = doc.add_paragraph()
                        spacing_run = spacing_para.add_run()
                        spacing_run.add_break()
                    
                    # Add signature paragraph
                    sig_para = doc.add_paragraph()
                    
                    # Set alignment
                    if is_right_aligned:
                        sig_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    else:
                        sig_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # Add signature label
                    run_label = sig_para.add_run("Signature")
                    run_label.bold = True
                    
                    # Add signature image
                    run_image = sig_para.add_run()
                    try:
                        image_width = Inches(2.0)
                        if is_right_aligned:
                            image_width = Inches(1.8)
                        run_image.add_picture(signature_path, width=image_width)
                    except Exception as e:
                        print(f"Error adding signature image: {e}")
                        run_image.add_run("[SIGNATURE IMAGE]")
                    
                    print(f"Successfully placed regular signature at document end")
            
            return True
            
        except Exception as e:
            print(f"Error placing regular signatures: {e}")
            import traceback
            traceback.print_exc()
            return False


class SignaturePlacementHelper:
    
    @staticmethod
    def suggest_signature_positions(text_content: str) -> Dict:
        """
        Analyze document text and suggest signature positions
        
        Args:
            text_content: Document text content
            
        Returns:
            Dictionary with suggested positions and reasoning
        """
        suggestions = {
            'positions': [],
            'confidence': 0.0,
            'reasoning': ''
        }
        
        # Keywords that typically precede signatures
        signature_keywords = [
            'signature', 'sign here', 'signed by', 'signature of',
            'please sign', 'witness signature', 'authorized signature'
        ]
        
        # Document type indicators
        contract_indicators = ['contract', 'agreement', 'terms', 'conditions']
        form_indicators = ['form', 'application', 'request']
        
        text_lower = text_content.lower()
        found_keywords = []
        
        # Look for signature-related keywords
        for keyword in signature_keywords:
            if keyword in text_lower:
                found_keywords.append(keyword)
        
        # Determine confidence based on findings
        if found_keywords:
            suggestions['confidence'] = min(0.9, 0.5 + len(found_keywords) * 0.1)
            suggestions['reasoning'] = f"Found signature keywords: {', '.join(found_keywords)}"
            
            # Suggest position at the end of document
            suggestions['positions'].append({
                'page': 1,
                'x': 1.0,  # 1 inch from left
                'y': 10.0,  # 10 inches from top (bottom of page)
                'reason': 'End of document placement'
            })
        else:
            # Fallback: suggest common position
            suggestions['confidence'] = 0.3
            suggestions['reasoning'] = "No signature keywords found, suggesting standard position"
            suggestions['positions'].append({
                'page': 1,
                'x': 1.0,
                'y': 10.0,
                'reason': 'Standard document end placement'
            })
        
        # Adjust for document type
        if any(indicator in text_lower for indicator in contract_indicators):
            suggestions['confidence'] = min(1.0, suggestions['confidence'] + 0.1)
            suggestions['reasoning'] += " - Contract document type detected"
        elif any(indicator in text_lower for indicator in form_indicators):
            suggestions['reasoning'] += " - Form document type detected"
        
        return suggestions